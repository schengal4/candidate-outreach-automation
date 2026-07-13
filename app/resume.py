"""Resume (.docx) text extraction and on-disk resume file handling."""

from io import BytesIO
from pathlib import Path

from docx import Document

from . import config
from .pdf_convert import docx_to_pdf

MAX_RESUME_CHARS = 30_000


def extract_docx_text(data: bytes) -> str:
    """Extract plain text from a .docx file (paragraphs + tables)."""
    doc = Document(BytesIO(data))
    parts = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    text = "\n".join(parts).strip()
    if not text:
        raise ValueError("No text could be extracted from the resume .docx file.")
    return text[:MAX_RESUME_CHARS]


def resume_docx_path(candidate) -> Path:
    return config.settings.RESUME_DIR / f"{candidate.id}_{candidate.resume_filename}"


def resume_pdf_path(candidate) -> Path:
    return config.settings.RESUME_DIR / f"{candidate.id}_{Path(candidate.resume_filename).stem}.pdf"


def resume_pdf_display_name(candidate) -> str:
    """Attachment filename the email recipient sees — the uploaded name, as PDF,
    without the internal candidate-id prefix."""
    return f"{Path(candidate.resume_filename).stem}.pdf"


def ensure_resume_pdf(candidate) -> Path:
    """Return the cached resume PDF, converting from the .docx if not yet cached.

    Normally the PDF is created once at upload time; the lazy conversion here
    covers candidates created before the attachment feature existed. Blocking
    when a conversion runs — wrap in asyncio.to_thread from the event loop.
    """
    pdf_path = resume_pdf_path(candidate)
    if pdf_path.exists():
        return pdf_path
    docx_path = resume_docx_path(candidate)
    if not docx_path.exists():
        raise FileNotFoundError(f"Resume file not found: {docx_path.name}")
    docx_to_pdf(docx_path, pdf_path)
    return pdf_path
